#!/usr/bin/env python3
"""
Test script for the Document Modifier
Tests full document modification capabilities with different file types
"""

import sys
import os

# Add the core directory to the path
sys.path.append(os.path.join(os.path.dirname(__file__), 'core'))

from detection_engine import PIIDetectionEngine
from document_modifier import DocumentModifier

def test_document_modifier():
    """Test the document modifier with different file types"""
    
    print("🧪 Testing Document Modifier")
    print("=" * 50)
    
    # Initialize the engine and modifier
    engine = PIIDetectionEngine()
    modifier = DocumentModifier(engine)
    
    print(f"Supported modification types: {modifier.get_supported_modification_types()}")
    
    # Test 1: Create and modify a CSV file
    print("\n📋 Test 1: CSV File Modification")
    print("-" * 30)
    
    test_csv_content = """Name,Email,Phone,Address
John Smith,john.smith@email.com,(555) 123-4567,123 Main St
Jane Doe,jane.doe@company.com,(555) 987-6543,456 Oak Ave
Bob Johnson,bob.j@business.net,(555) 456-7890,789 Pine Rd"""
    
    with open('test_modify.csv', 'w', newline='') as f:
        f.write(test_csv_content)
    
    try:
        result = modifier.modify_file('test_modify.csv', mask_format="token")
        
        print(f"Original: {result.original_file}")
        print(f"Modified: {result.modified_file}")
        print(f"Type: {result.file_type}")
        print(f"Entities masked: {result.entities_masked}")
        print(f"Processing time: {result.processing_time:.2f} seconds")
        
        if result.errors:
            print(f"Errors: {result.errors}")
        else:
            print("✅ CSV modification successful!")
            
            # Show the modified content
            with open(result.modified_file, 'r') as f:
                modified_content = f.read()
            print("\nModified content:")
            print(modified_content)
        
    except Exception as e:
        print(f"Error modifying CSV: {e}")
    
    finally:
        # Clean up
        if os.path.exists('test_modify.csv'):
            os.remove('test_modify.csv')
        if os.path.exists('test_modify_masked.csv'):
            os.remove('test_modify_masked.csv')
    
    # Test 2: Create and modify a text file
    print("\n📋 Test 2: Text File Modification")
    print("-" * 30)
    
    test_text_content = """
    CONFIDENTIAL DOCUMENT
    
    Client Information:
    - Name: John Smith
    - Email: john.smith@company.com
    - Phone: (555) 987-6543
    - Address: 456 Oak Avenue, Somewhere, CA 90210
    - SSN: 987-65-4321
    - Credit Card: 5555-4444-3333-2222
    - IP Address: 10.0.0.1
    
    Project Details: This project involves sensitive data processing.
    """
    
    with open('test_modify.txt', 'w') as f:
        f.write(test_text_content)
    
    try:
        result = modifier.modify_file('test_modify.txt', mask_format="token")
        
        print(f"Original: {result.original_file}")
        print(f"Modified: {result.modified_file}")
        print(f"Type: {result.file_type}")
        print(f"Entities masked: {result.entities_masked}")
        print(f"Processing time: {result.processing_time:.2f} seconds")
        
        if result.errors:
            print(f"Errors: {result.errors}")
        else:
            print("✅ Text modification successful!")
            
            # Show the modified content
            with open(result.modified_file, 'r') as f:
                modified_content = f.read()
            print("\nModified content:")
            print(modified_content)
        
    except Exception as e:
        print(f"Error modifying text: {e}")
    
    finally:
        # Clean up
        if os.path.exists('test_modify.txt'):
            os.remove('test_modify.txt')
        if os.path.exists('test_modify_masked.txt'):
            os.remove('test_modify_masked.txt')
    
    # Test 3: Create and modify a DOCX file
    print("\n📋 Test 3: DOCX File Modification")
    print("-" * 30)
    
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Document', 0)
        
        p = doc.add_paragraph('This is a test document with PII data.')
        p.add_run('\nContact: John Smith (john.smith@company.com)')
        p.add_run('\nPhone: (555) 123-4567')
        p.add_run('\nSSN: 123-45-6789')
        
        doc.save('test_modify.docx')
        
        result = modifier.modify_file('test_modify.docx', mask_format="token")
        
        print(f"Original: {result.original_file}")
        print(f"Modified: {result.modified_file}")
        print(f"Type: {result.file_type}")
        print(f"Entities masked: {result.entities_masked}")
        print(f"Processing time: {result.processing_time:.2f} seconds")
        
        if result.errors:
            print(f"Errors: {result.errors}")
        else:
            print("✅ DOCX modification successful!")
        
    except Exception as e:
        print(f"Error modifying DOCX: {e}")
    
    finally:
        # Clean up
        if os.path.exists('test_modify.docx'):
            os.remove('test_modify.docx')
        if os.path.exists('test_modify_masked.docx'):
            os.remove('test_modify_masked.docx')
    
    # Test 4: Create and modify a PDF file
    print("\n📋 Test 4: PDF File Modification")
    print("-" * 30)
    
    try:
        import fitz  # PyMuPDF
        
        # Create a simple PDF with test content
        doc = fitz.open()
        page = doc.new_page()
        
        text_content = """
        CONFIDENTIAL REPORT
        
        Client Information:
        - Name: John Smith
        - Email: john.smith@company.com
        - Phone: (555) 123-4567
        - Address: 123 Main Street, Anytown, ST 12345
        - SSN: 123-45-6789
        - Credit Card: 4111-1111-1111-1111
        - IP Address: 192.168.1.1
        
        Project Details:
        This confidential report contains sensitive information that must be protected.
        All PII should be redacted before sharing with external parties.
        """
        
        page.insert_text((50, 50), text_content)
        doc.save('test_modify.pdf')
        doc.close()
        
        result = modifier.modify_file('test_modify.pdf', mask_format="token")
        
        print(f"Original: {result.original_file}")
        print(f"Modified: {result.modified_file}")
        print(f"Type: {result.file_type}")
        print(f"Entities masked: {result.entities_masked}")
        print(f"Processing time: {result.processing_time:.2f} seconds")
        
        if result.errors:
            print(f"Errors: {result.errors}")
        else:
            print("✅ PDF modification successful!")
        
    except Exception as e:
        print(f"Error modifying PDF: {e}")
    
    finally:
        # Clean up
        if os.path.exists('test_modify.pdf'):
            os.remove('test_modify.pdf')
        if os.path.exists('test_modify_masked.pdf'):
            os.remove('test_modify_masked.pdf')
    
    print("\n✅ Document modifier test completed!")

if __name__ == "__main__":
    test_document_modifier()
