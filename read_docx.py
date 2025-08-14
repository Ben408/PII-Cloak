#!/usr/bin/env python3
"""
Script to extract text content from .docx files
"""

import sys
from docx import Document
import os

def read_docx(file_path):
    """Extract text from a .docx file"""
    try:
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

def main():
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <docx_file>")
        print("Available .docx files in current directory:")
        for file in os.listdir('.'):
            if file.endswith('.docx') and not file.startswith('~$'):
                print(f"  - {file}")
        return
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    if not file_path.endswith('.docx'):
        print("Please provide a .docx file")
        return
    
    print(f"Reading content from: {file_path}")
    print("=" * 50)
    
    content = read_docx(file_path)
    print(content)

if __name__ == "__main__":
    main()
